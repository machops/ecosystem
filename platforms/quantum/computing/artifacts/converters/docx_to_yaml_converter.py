# @ECO-layer: GQS-L0
# @ECO-governed
# @ECO-layer: GL00-09
# @ECO-semantic: general-component
# @ECO-audit-trail: gl-enterprise-architecture/gl-platform.governance/audit-trails/GL00_09-audit.json
#
# GL Unified Architecture Governance Framework Activated
# GL Root Semantic Anchor: gl-enterprise-architecture/gl-platform.governance/ECO-ROOT-SEMANTIC-ANCHOR.yaml
# GL Unified Naming Charter: gl-enterprise-architecture/gl-platform.governance/ECO-UNIFIED-NAMING-CHARTER.yaml


#!/usr/bin/env python3
"""
GL Quantum Artifact Converter - DOCX to YAML
Converts DOCX documents to structured YAML artifacts
"""

"""
Module docstring
================

This module is part of the GL governance framework.
Please add specific module documentation here.
"""
# MNGA-002: Import organization needs review
import os
import sys
import yaml
import json
from pathlib import Path
from typing import Dict, Any, List
from docx import Document
from datetime import datetime

class DOCXToYAMLConverter:
    """Convert DOCX documents to structured YAML artifacts"""
    
    def __init__(self, input_path: str, output_path: str = None):
        self.input_path = Path(input_path)
        self.output_path = Path(output_path) if output_path else self.input_path.with_suffix('.yaml')
        
    def extract_document_structure(self, doc: Document) -> Dict[str, Any]:
        """Extract structured content from DOCX"""
        
        structure = {
            'metadata': {
                'title': '',
                'author': '',
                'created': datetime.now().isoformat(),
                'version': '1.0.0',
                'type': 'artifact',
                'format': 'yaml',
                'source': str(self.input_path)
            },
            'content': {
                'sections': [],
                'tables': [],
                'images': [],
                'metadata': {}
            },
            'validation': {
                'required_fields': True,
                'structure_valid': True
            }
        }
        
        # Extract paragraphs with hierarchy
        current_section = None
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                if current_section:
                    structure['content']['sections'].append(current_section)
                current_section = {
                    'level': int(para.style.name.split(' ')[-1]),
                    'heading': para.text,
                    'content': []
                }
            elif current_section:
                current_section['content'].append(para.text)
            elif para.text.strip():
                # Document metadata or intro
                if not structure['metadata']['title']:
                    structure['metadata']['title'] = para.text
                    
        if current_section:
            structure['content']['sections'].append(current_section)
            
        # Extract tables
        for table in doc.tables:
            table_data = {
                'rows': len(table.rows),
                'columns': len(table.columns),
                'headers': [cell.text for cell in table.rows[0].cells],
                'data': []
            }
            for row in table.rows[1:]:
                row_data = [cell.text for cell in row.cells]
                table_data['data'].append(row_data)
            structure['content']['tables'].append(table_data)
            
        # Extract document properties
        core_props = doc.core_properties
        structure['metadata']['author'] = core_props.author or 'Unknown'
        structure['metadata']['created'] = core_props.created.isoformat() if core_props.created else datetime.now().isoformat()
        structure['metadata']['modified'] = core_props.modified.isoformat() if core_props.modified else datetime.now().isoformat()
        
        return structure
        
    def convert(self) -> Dict[str, Any]:
        """Convert DOCX to YAML structure"""
        
        if not self.input_path.exists():
            raise FileNotFoundError(f"Input file not found: {self.input_path}")
            
        doc = Document(str(self.input_path))
        structure = self.extract_document_structure(doc)
        
        return structure
        
    def save_yaml(self, data: Dict[str, Any], output_path: Path = None):
        """Save YAML artifact"""
        
        output = output_path or self.output_path
        output.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output, 'w', encoding='utf-8') as f:
            yaml.dump(data, f, default_flow_style=False, sort_keys=False, allow_unicode=True)
            
        return output
        
    def validate_structure(self, data: Dict[str, Any]) -> bool:
        """Validate YAML structure"""
        
        required_keys = ['metadata', 'content', 'validation']
        for key in required_keys:
            if key not in data:
                return False
                
        return True

def main():
    """Main entry point"""
    
    if len(sys.argv) < 2:
        print("Usage: python docx-to-yaml-converter.py <input.docx> [output.yaml]")
        sys.exit(1)
        
    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    converter = DOCXToYAMLConverter(input_path, output_path)
    
    try:
        structure = converter.convert()
        
        if converter.validate_structure(structure):
            output_file = converter.save_yaml(structure)
            print(f"✅ Successfully converted {input_path} to {output_file}")
            return 0
        else:
            print("❌ Validation failed: Invalid structure")
            return 1
            
    except Exception as e:
        print(f"❌ Error: {e}")
        return 1

if __name__ == '__main__':
    sys.exit(main())