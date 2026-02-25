# @ECO-governed
# @ECO-layer: GL00-09
# @ECO-semantic: general-component
# @ECO-audit-trail: gl-enterprise-architecture/gl_platform_universe.governance/audit-trails/GL00_09-audit.json
#
# GL Unified Charter Activated
# GL Root Semantic Anchor: gl-enterprise-architecture/gl_platform_universe.governance/ECO-ROOT-SEMANTIC-ANCHOR.yaml
# GL Unified Naming Charter: gl-enterprise-architecture/gl_platform_universe.governance/ECO-UNIFIED-NAMING-CHARTER.yaml


#!/usr/bin/env python3
"""
Convert DOCX documents to structured artifacts (YAML/JSON/Markdown)
"""

"""
Module docstring
================

This module is part of the GL governance framework.
Please add specific module documentation here.
"""
# MNGA-002: Import organization needs review
import sys
import json
import yaml
from pathlib import Path
from docx import Document
from markdownify import markdownify as md

def extract_docx_content(docx_path):
    """Extract content from DOCX file"""
    doc = Document(docx_path)
    
    content = {
        'title': '',
        'sections': [],
        'tables': [],
        'metadata': {}
    }
    
    # Extract paragraphs and headings
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading 1'):
            content['title'] = paragraph.text
        elif paragraph.style.name.startswith('Heading'):
            content['sections'].append({
                'level': int(paragraph.style.name.split()[-1]),
                'text': paragraph.text
            })
        elif paragraph.text.strip():
            content['sections'].append({
                'level': 0,
                'text': paragraph.text
            })
    
    # Extract tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        content['tables'].append(table_data)
    
    # Extract metadata from document properties
    content['metadata'] = {
        'author': doc.core_properties.author,
        'created': str(doc.core_properties.created) if doc.core_properties.created else None,
        'modified': str(doc.core_properties.modified) if doc.core_properties.modified else None,
        'title': doc.core_properties.title or content['title']
    }
    
    return content

def convert_to_yaml(content, output_path):
    """Convert content to YAML"""
    yaml_content = yaml.dump(content, default_flow_style=False, sort_keys=False)
    
    output_file = Path(output_path) / 'artifact.yaml'
    with open(output_file, 'w') as f:
        f.write(yaml_content)
    
    return output_file

def convert_to_json(content, output_path):
    """Convert content to JSON"""
    json_content = json.dumps(content, indent=2)
    
    output_file = Path(output_path) / 'artifact.json'
    with open(output_file, 'w') as f:
        f.write(json_content)
    
    return output_file

def convert_to_markdown(content, output_path):
    """Convert content to Markdown"""
    md_content = f"# {content['title']}\n\n"
    
    for section in content['sections']:
        if section['level'] > 0:
            md_content += f"{'#' * (section['level'] + 1)} {section['text']}\n\n"
        else:
            md_content += f"{section['text']}\n\n"
    
    for table in content['tables']:
        md_content += "| " + " | ".join(table[0]) + " |\n"
        md_content += "| " + " | ".join(["---"] * len(table[0])) + " |\n"
        for row in table[1:]:
            md_content += "| " + " | ".join(row) + " |\n"
        md_content += "\n"
    
    output_file = Path(output_path) / 'artifact.md'
    with open(output_file, 'w') as f:
        f.write(md_content)
    
    return output_file

def main():
    if len(sys.argv) < 2:
        print("Usage: convert.py <source-path> [output-path] [format]", file=sys.stderr)
        sys.exit(1)
    
    source_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else 'artifacts/modules'
    format = sys.argv[3] if len(sys.argv) > 3 else 'yaml'
    
    # Create output directory
    Path(output_path).mkdir(parents=True, exist_ok=True)
    
    # Extract content
    content = extract_docx_content(source_path)
    
    # Convert based on format
    if format == 'yaml':
        output_file = convert_to_yaml(content, output_path)
    elif format == 'json':
        output_file = convert_to_json(content, output_path)
    elif format == 'markdown':
        output_file = convert_to_markdown(content, output_path)
    else:
        print(f"Unsupported format: {format}", file=sys.stderr)
        sys.exit(1)
    
    print(f"✅ Successfully converted to {format}: {output_file}")

if __name__ == '__main__':
    main()
