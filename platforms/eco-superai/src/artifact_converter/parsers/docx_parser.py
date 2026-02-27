"""DOCX parser with graceful fallback.

Uses ``python-docx`` when available; otherwise provides a ZIP-based fallback
that extracts raw text from the XML within the DOCX container.
"""

from __future__ import annotations

import io
import re
import zipfile
from pathlib import Path
from typing import Any

import structlog

from . import BaseParser, ParseResult

logger = structlog.get_logger(__name__)

try:
    import docx as python_docx  # type: ignore[import-untyped]

    _HAS_PYTHON_DOCX = True
except ImportError:  # pragma: no cover
    _HAS_PYTHON_DOCX = False
    logger.info(
        "python_docx_not_installed",
        hint="Install python-docx for full DOCX support: pip install python-docx",
    )


class DocxParser(BaseParser):
    """Parser for Microsoft Word ``.docx`` files."""

    def parse(self, content: str | bytes, source_path: Path | None = None) -> ParseResult:
        if isinstance(content, str):
            logger.warning("docx_received_text_content", source=str(source_path))
            return ParseResult(body=content, raw=content)

        if _HAS_PYTHON_DOCX:
            return self._parse_with_python_docx(content, source_path)

        return self._parse_fallback(content, source_path)

    def supported_extensions(self) -> list[str]:
        return [".docx"]

    # ------------------------------------------------------------------
    # python-docx-based extraction
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_with_python_docx(
        content: bytes, source_path: Path | None
    ) -> ParseResult:
        logger.debug("docx_python_docx_start", source=str(source_path))

        doc = python_docx.Document(io.BytesIO(content))

        metadata: dict[str, Any] = {}
        core = doc.core_properties
        if core.title:
            metadata["title"] = core.title
        if core.author:
            metadata["author"] = core.author
        if core.created:
            metadata["date"] = core.created.isoformat()
        if core.keywords:
            metadata["tags"] = core.keywords
        if core.subject:
            metadata["description"] = core.subject

        sections: list[dict[str, Any]] = []
        current_heading: str = ""
        current_content_lines: list[str] = []

        for paragraph in doc.paragraphs:
            style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
            text = paragraph.text.strip()
            if not text:
                continue

            if "heading" in style_name:
                # Flush previous section
                if current_heading or current_content_lines:
                    sections.append({
                        "heading": current_heading,
                        "content": "\n".join(current_content_lines).strip(),
                    })
                current_heading = text
                current_content_lines = []
            else:
                current_content_lines.append(text)

        # Flush final section
        if current_heading or current_content_lines:
            sections.append({
                "heading": current_heading,
                "content": "\n".join(current_content_lines).strip(),
            })

        body = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

        logger.debug(
            "docx_python_docx_done",
            paragraphs=len(doc.paragraphs),
            sections=len(sections),
        )

        return ParseResult(
            body=body.strip(),
            metadata=metadata,
            sections=sections,
            raw=body,
        )

    # ------------------------------------------------------------------
    # ZIP/XML fallback
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_fallback(content: bytes, source_path: Path | None) -> ParseResult:
        logger.warning(
            "docx_fallback_extraction",
            hint="Install python-docx for full DOCX support.",
            source=str(source_path),
        )

        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                if "word/document.xml" not in zf.namelist():
                    logger.error("docx_missing_document_xml", source=str(source_path))
                    return ParseResult(
                        body="",
                        metadata={"_fallback": True, "_error": "Not a valid DOCX file"},
                    )

                xml_content = zf.read("word/document.xml").decode("utf-8", errors="replace")

            # Strip XML tags to get raw text
            text = re.sub(r"<[^>]+>", " ", xml_content)
            text = re.sub(r"\s+", " ", text).strip()

            return ParseResult(
                body=text,
                metadata={
                    "_fallback": True,
                    "_warning": "python-docx not installed; structure may be lost",
                },
                sections=[{"heading": "", "content": text}] if text else [],
                raw=text,
            )
        except zipfile.BadZipFile:
            logger.error("docx_bad_zip", source=str(source_path))
            return ParseResult(
                body="",
                metadata={"_fallback": True, "_error": "Invalid DOCX (bad ZIP)"},
            )
